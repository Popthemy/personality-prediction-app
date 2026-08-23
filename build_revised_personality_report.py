from __future__ import annotations
import json
import sqlite3
from pathlib import Path
from statistics import mean

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\USER\Demola\personality-prediction-app")
SOURCE_DOCX = Path(r"C:\Users\USER\Downloads\Personality Prediction 1-3 .docx")
DB_PATH = ROOT / "db.sqlite3"
OUTPUT_DOCX = ROOT / "revised_personality_prediction_full_report.docx"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def style_body_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.space_after = Pt(after)
    pf.line_spacing = line
    return paragraph


def style_heading(paragraph, level=1):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6)
    if level == 1:
        for run in paragraph.runs:
            set_run_font(run, size=14, bold=True)
    elif level == 2:
        for run in paragraph.runs:
            set_run_font(run, size=12, bold=True)
    else:
        for run in paragraph.runs:
            set_run_font(run, size=12, bold=True)
    return paragraph


def add_body_paragraph(doc, text, bold_prefix=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    style_body_paragraph(p, align=align)
    if bold_prefix and text.startswith(bold_prefix):
        prefix, rest = text.split(":", 1)
        run = p.add_run(prefix + ":")
        set_run_font(run, bold=True)
        run2 = p.add_run(rest)
        set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return style_heading(p, level=level)


def add_table(doc, rows, col_widths=None, header_fill=None):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    if col_widths:
        for idx, width in enumerate(col_widths):
            table.add_column(Inches(width)) if False else None
    hdr = table.add_row().cells
    for j, value in enumerate(rows[0]):
        hdr[j].text = str(value)
        for run in hdr[j].paragraphs[0].runs:
            set_run_font(run, size=10, bold=True)
        hdr[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
            for run in cells[j].paragraphs[0].runs:
                set_run_font(run, size=10)
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if j else WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
    return table


def replace_paragraph_contains(doc, needle, new_text):
    for p in doc.paragraphs:
        if needle in p.text:
            p.text = new_text
            for run in p.runs:
                set_run_font(run)
            style_body_paragraph(p)
            return True
    return False


def find_paragraph_index(doc, needle):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle or needle in p.text:
            return i
    return None


def replace_abstract(doc):
    idx = find_paragraph_index(doc, "ABSTRACT")
    if idx is None:
        return
    abstract_paragraphs = [
        "This study implements a web-based personality prediction system that combines psychometric assessment, social media text processing, contextual language representation, embedding-space augmentation, and sparse regression to infer the Big Five personality traits. The delivered application is a Django Model-View-Template system with researcher accounts, BFI-44 survey import, post acquisition, machine-learning orchestration, dashboard analytics, and a public-facing live prediction demo.",
        "The implementation does not rely on a theoretical-only workflow. Instead, it stores volunteer records, BFI-44 responses, X timeline posts, BERT embeddings, synthetic augmentation samples, Lasso models, and final psychometric profiles in a relational database. Where the written design originally described a fully adversarial GAN and a fully trained reinforcement-learning policy, the actual codebase implements a simplified augmentation module and a Q-learning style selection scaffold that support the overall pipeline without overstating what is currently automated.",
        "Evaluation of the local dataset shows that the system is functional end to end, but also reveals the realistic limitations of a small academic prototype. Across the currently stored profiles, the mean overall MAE is 0.474, the mean prediction confidence is 0.772, and the average number of posts analyzed per profile is 8.25. These results indicate that the platform is usable and traceable, while also showing that model generalization still depends strongly on more data and stronger validation.",
        "The report therefore presents the implemented system as it actually exists, explains the corrections made to the earlier write-up, and discusses the practical and methodological lessons that follow from the deployment of the prototype in a real project setting.",
    ]
    for offset, text in enumerate(abstract_paragraphs, start=1):
        p = doc.paragraphs[idx + offset]
        p.text = text
        for run in p.runs:
            set_run_font(run)
        style_body_paragraph(p)


def patch_existing_claims(doc):
    replacements = [
        (
            "The proposed methodology adopts a sequential architecture in which Q-Learning is initially employed for adaptive feature selection.",
            "The implemented pipeline follows the same broad sequence, but the deployed code uses a practical post selection pass, BERT encoding, simplified augmentation, and regularized regression rather than a purely theoretical reinforcement-learning training loop."
        ),
        (
            "To mitigate challenges related to limited and imbalanced datasets, Generative Adversarial Networks (GANs) are incorporated to perform data augmentation through the synthesis of realistic text representations, enhancing the diversity and robustness of the training data.",
            "To mitigate the effect of limited data, the codebase applies embedding-space perturbation with template-based synthetic text generation. This is presented in the report as a GAN-inspired augmentation module, but the actual implementation is a simplified augmenter rather than a full discriminator-generator adversarial network."
        ),
        (
            "The present research systematically develops and evaluates an integrated machine learning pipeline specifically designed to predict the Big Five (OCEAN) personality traits from social media-style textual and profile data.",
            "The present research systematically develops and evaluates an integrated machine learning pipeline for predicting the Big Five (OCEAN) personality traits from social media-style textual and profile data within a Django-based application."
        ),
        (
            "The methodological pipeline operates in strict sequential order: Q-Learning active selection then BERT contextual encoding then GAN augmentation (applied exclusively during training) then Lasso sparse prediction. Evaluation employs 5-fold stratified cross-validation to produce continuous OCEAN scores normalized to a 0–1 interval.",
            "The implementation follows a sequential order of data acquisition, text cleaning, Q-learning based post selection, BERT contextual encoding, augmentation, and Lasso or ElasticNet regression. In practice, the stored runs rely on persisted database artifacts and cohort-based training summaries rather than a purely abstract cross-validation description."
        ),
        (
            "The research adopts a quantitative, model-driven, and experimental design under the umbrella of Design Science Research (DSR).",
            "The research adopts a quantitative, model-driven, and experimental design under the umbrella of Design Science Research (DSR), with the application serving as a working software artifact that can be inspected, tested, and iteratively improved."
        ),
        (
            "The implementation follows a staged progression:\nPhase 1: Data harvesting and ground-truth labeling (X API and BFI-44).\nPhase 2: Reinforcement learning-based active signal selection (Q-Learning).\nPhase 3: Contextual semantic extraction via transformer encoders (BERT).\nPhase 4: Latent manifold expansion using generative adversarial networks (GAN).",
            "The implementation follows a staged progression: Phase 1 data acquisition and ground-truth labeling through CSV import of BFI-44 responses and X timeline retrieval, Phase 2 text cleaning and post selection, Phase 3 contextual semantic extraction via transformer encoders, Phase 4 simplified embedding-space augmentation, and Phase 5 sparse regression and psychometric profile generation."
        ),
        (
            "The current study addresses identified bottlenecks by introducing an intelligent, integrated pipeline that processes natural language directly from X (Twitter). The architecture is built on a modular MVT (Model-View-Template) pattern, utilizing Django for the web interface and Celery for asynchronous processing.",
            "The current study addresses identified bottlenecks by introducing an intelligent, integrated pipeline that processes social media text directly from X timelines and imported datasets. The architecture is built on a modular MVT (Model-View-Template) pattern, utilizing Django for the web interface and Celery and Redis support for asynchronous processing."
        ),
        (
            "All experiments are executed on Google Colab Pro infrastructure equipped with T4 and A100 GPU acceleration.",
            "The implementation is realized as a Django web application backed by SQLite in development, PostgreSQL-ready models for deployment, and a local machine-learning workflow that can use GPU acceleration when BERT is loaded on compatible hardware."
        ),
        (
            "Real-time integration with social media APIs is deliberately excluded to ensure full compliance with privacy regulations and platform policies.",
            "Real-time acquisition is implemented through a local X timeline fetcher using the xtf library and Nitter fallbacks, while the BFI-44 ground truth is imported from CSV submissions."
        ),
        (
            "manual implementation of sparse linear regression",
            "scikit-learn-based Lasso and ElasticNet regression with feature scaling, label normalization, and stored model metadata"
        ),
        (
            "first principles manual implementation of sparse linear regression",
            "scikit-learn-based sparse regression with explicit scaling, coefficient persistence, and validation metrics"
        ),
    ]
    for old, new in replacements:
        replace_paragraph_contains(doc, old, new)


def query_db_stats():
    conn = sqlite3.connect(str(DB_PATH))
    cur = conn.cursor()
    stats = {}
    for table in [
        "volunteer",
        "bfi_survey",
        "post",
        "bert_embedding",
        "q_learning_log",
        "synthetic_data",
        "lasso_model",
        "psychometric_profile",
    ]:
        cur.execute(f"SELECT COUNT(*) FROM {table}")
        stats[table] = cur.fetchone()[0]

    cur.execute(
        "SELECT overall_mae, correlation, r2_score, posts_analyzed, embeddings_used, synthetic_data_used, prediction_confidence "
        "FROM psychometric_profile"
    )
    rows = cur.fetchall()
    if rows:
        stats["profile_means"] = {
            "overall_mae": mean(r[0] for r in rows if r[0] is not None),
            "correlation": mean(r[1] for r in rows if r[1] is not None),
            "r2_score": mean(r[2] for r in rows if r[2] is not None),
            "posts_analyzed": mean(r[3] for r in rows if r[3] is not None),
            "embeddings_used": mean(r[4] for r in rows if r[4] is not None),
            "synthetic_data_used": mean(r[5] for r in rows if r[5] is not None),
            "prediction_confidence": mean(r[6] for r in rows if r[6] is not None),
        }
        stats["profile_minmax"] = {
            "overall_mae": (min(r[0] for r in rows if r[0] is not None), max(r[0] for r in rows if r[0] is not None)),
            "correlation": (min(r[1] for r in rows if r[1] is not None), max(r[1] for r in rows if r[1] is not None)),
            "r2_score": (min(r[2] for r in rows if r[2] is not None), max(r[2] for r in rows if r[2] is not None)),
        }
    else:
        stats["profile_means"] = {}
        stats["profile_minmax"] = {}

    cur.execute("SELECT COUNT(*) FROM post WHERE selected_by_qlearning=1")
    stats["selected_posts"] = cur.fetchone()[0]
    cur.execute("SELECT COUNT(*) FROM post WHERE embedding_processed=1")
    stats["embedded_posts"] = cur.fetchone()[0]
    cur.execute("SELECT COUNT(*) FROM post WHERE q_value IS NOT NULL")
    stats["qvalued_posts"] = cur.fetchone()[0]

    cur.execute(
        "SELECT trait, AVG(train_mae), AVG(train_r2), AVG(training_samples_used), AVG(synthetic_samples_used) "
        "FROM lasso_model GROUP BY trait"
    )
    stats["lasso_trait_summary"] = cur.fetchall()

    cur.execute(
        "SELECT v.x_handle, p.overall_mae, p.correlation, p.r2_score, p.posts_analyzed, "
        "p.synthetic_data_used, p.prediction_confidence "
        "FROM psychometric_profile p JOIN volunteer v ON p.volunteer_id = v.id "
        "ORDER BY p.created_at DESC"
    )
    stats["profile_rows"] = cur.fetchall()
    conn.close()
    return stats


def chapter_four(doc, stats):
    doc.add_page_break()
    add_heading(doc, "CHAPTER FOUR", 1)
    add_heading(doc, "SYSTEM IMPLEMENTATION, TESTING, RESULTS AND DISCUSSION", 1)

    add_heading(doc, "4.1 Introduction", 2)
    add_body_paragraph(
        doc,
        "This chapter describes the system as it is implemented in the repository, not merely as it was imagined in the proposal stage. The application is a Django-based personality prediction platform that combines researcher accounts, BFI-44 import, timeline retrieval, preprocessing, BERT embeddings, simplified augmentation, regularized regression, and dashboard reporting. The aim of the chapter is to show the relationship between the code, the database, and the reported analysis so that the final write-up remains technically honest."
    )
    add_body_paragraph(
        doc,
        "The earlier write-up described a more idealized machine-learning stack than the code currently executes in routine runs. Some of those ideas are still present as design goals, but the implementation is more pragmatic. For that reason, this chapter documents the actual runtime behavior of the platform, highlights the parts that are fully implemented, and also points out where the current code remains a simplified academic prototype rather than a fully generalized production model."
    )

    add_heading(doc, "4.2 Alignment Between the Written Design and the Implemented System", 2)
    add_body_paragraph(
        doc,
        "A useful way to understand the final system is to compare the written design with the code that was produced. This comparison is important because the project evolved from a conceptual multi-paradigm architecture into a working web application with a smaller but more concrete set of behaviors. The table below summarizes the main alignments and corrections."
    )
    alignment_rows = [
        ["Report Claim", "Actual Code Behavior", "Adjustment Reflected in This Report"],
        ["Q-learning active selection", "QLearningAgent exists, but the full pipeline currently performs greedy top-k selection and does not persist routine episode logs.", "Describe the module as an active selection scaffold with a practical post-ranking implementation."],
        ["BERT embeddings", "Fully implemented with bert-base-uncased CLS embeddings, GPU fallback, and database persistence.", "Retain this as a fully implemented component."],
        ["GAN augmentation", "Implemented as Gaussian perturbation of embeddings plus template-based text generation.", "Describe it as simplified augmentation rather than a true adversarial GAN."],
        ["Manual Lasso from scratch", "Implemented with scikit-learn Lasso/ElasticNet, scaling, cross-validation, and stored coefficients.", "Replace the manual-from-scratch claim with the actual library-backed implementation."],
        ["X API data collection", "Implemented through xtf/Nitter fetch fallbacks and CSV import for BFI-44 ground truth.", "State the actual acquisition path and fallback logic."],
    ]
    add_table(doc, alignment_rows, col_widths=[1.7, 2.6, 2.2])
    add_body_paragraph(
        doc,
        "This alignment exercise is not a weakness in the project. On the contrary, it is a healthy correction that many final-year projects benefit from. A report that states only the intended design but not the delivered code will always create problems during presentation and defense. The revised chapter therefore gives priority to implementation truth, while still acknowledging the conceptual direction that informed the design."
    )

    add_heading(doc, "4.3 Hardware and Software Environment", 2)
    hardware_rows = [
        ["Component", "Implemented Environment"],
        ["Processor", "A standard Intel Core i5 class system or higher is sufficient for routine operation"],
        ["Memory", "At least 8 GB RAM for normal development; more is recommended for BERT loading"],
        ["Storage", "SSD storage for the database, logs, and exported cleaned timelines"],
        ["Operating System", "Windows 10/11 or Ubuntu Linux"],
        ["Internet Access", "Required for timeline retrieval and model downloads when first loading BERT"],
    ]
    add_table(doc, hardware_rows, col_widths=[1.8, 4.7])
    add_body_paragraph(
        doc,
        "The software environment is anchored on Django for the web layer and Python for the machine-learning services. The repository includes the backend application, template views, model classes, form handling, and service modules that support preprocessing, embedding, augmentation, and regression. The stack also includes scikit-learn for regularized regression, Hugging Face Transformers and PyTorch for BERT loading, Celery and Redis for asynchronous execution support, and Chart.js on the frontend for the radar-style personality display."
    )
    add_body_paragraph(
        doc,
        "For data acquisition, the code relies on the xtf library and Nitter-based timeline retrieval rather than on a direct official X API key workflow. This is a practical choice for a student project because it allows public timeline acquisition to continue even when a particular instance becomes unavailable. The BFI-44 questionnaire is imported from CSV exports, which makes ground-truth labeling easier to manage and repeat during testing."
    )

    add_heading(doc, "4.4 Data Acquisition, Cleaning and Ground-Truth Labeling", 2)
    add_body_paragraph(
        doc,
        "The project supports two complementary forms of input. The first is the BFI-44 ground-truth survey, which arrives as a CSV export from the questionnaire workflow. The CSV uploader extracts the volunteer handle, verifies consent, parses the 44 item responses, and computes OCEAN scores using the BFIScorer service. The second is the social media timeline, which is stored as POST records and may be obtained either from the database or from the live fetch fallback when the local table is empty."
    )
    add_body_paragraph(
        doc,
        "The scoring logic for the BFI-44 instrument is important because it creates the labels used for model training and evaluation. The service performs the reverse-scoring of negatively worded items, validates that the response set is complete enough for scoring, and calculates openness, conscientiousness, extraversion, agreeableness, and neuroticism on the standard 1-to-5 scale. This ensures that the regression models are trained against meaningful psychometric targets rather than arbitrary placeholder values."
    )
    add_body_paragraph(
        doc,
        "Timeline cleaning is handled separately by the text preprocessing pipeline. The cleaner removes URLs, mentions, retweet markers, leading punctuation, and excessive whitespace. Hashtags are normalized so that the useful lexical content is retained, while non-ASCII symbols are stripped by default. Posts shorter than the configured minimum length are excluded so that the downstream embedding stage receives text that has at least some personality-bearing substance."
    )
    add_body_paragraph(
        doc,
        "This cleaning step matters because personality inference is sensitive to noise. A tweet that consists only of a mention, a link, or a short reaction does not carry the same linguistic signal as a longer self-expressive post. By filtering and normalizing the text before BERT encoding, the pipeline improves the chance that the resulting embeddings reflect actual expressive content rather than platform clutter."
    )

    add_heading(doc, "4.5 Core Module Implementation", 2)
    add_heading(doc, "4.5.1 User Registration and Researcher Workflow", 3)
    add_body_paragraph(
        doc,
        "The accounts application provides the authentication layer used by the researcher. It enables registration, login, and profile management, and it links each volunteer to the researcher who owns the record. In the context of this project, that relationship is important because it keeps the imported BFI survey, fetched timeline posts, and generated profile aligned with the correct study participant."
    )
    add_body_paragraph(
        doc,
        "The dashboard also gives the researcher a compact view of how many volunteers have been registered, how many already have BFI surveys, and how many have completed a prediction run. This is more than a cosmetic feature. It is the operational control center that turns the machine-learning pipeline into a usable study management system rather than a disconnected script."
    )

    add_heading(doc, "4.5.2 BFI-44 Survey Import and Scoring", 3)
    add_body_paragraph(
        doc,
        "The BFI upload workflow is one of the clearest examples of how the written report and the implementation now meet in the middle. The user uploads a CSV file, the code checks consent, identifies the X handle, extracts the item responses, and computes the Big Five scores. The scoring is written back into the database using the BFI_SURVEY model, and the volunteer object is updated so that later pipeline stages can confirm that a labeled ground truth exists."
    )
    add_body_paragraph(
        doc,
        "This design is robust because it supports re-uploading of corrected survey files without crashing the application. The upsert behavior means that a volunteer can be reprocessed when the ground truth needs to be refreshed, which is an important requirement in a real research setting where data cleaning often happens after initial collection."
    )

    add_heading(doc, "4.5.3 Timeline Retrieval and Post Persistence", 3)
    add_body_paragraph(
        doc,
        "Timeline retrieval is implemented by the TwitterFetcher service. The fetcher first tries the shared xtf router, then falls back to individual Nitter instances, and finally attempts a direct profile-page read if the earlier paths do not return posts. This layered approach is valuable because it makes the system resilient to the kind of retrieval failures that often affect public social media endpoints."
    )
    add_body_paragraph(
        doc,
        "When posts are successfully obtained, the fetcher deduplicates them using the X post identifier, computes engagement values, marks whether a post is a reply or retweet, and saves the results to the POST model. The pipeline then cleans the stored text and exports a sanitized timeline file for auditability. This means that every later stage can be traced back to a concrete text source instead of to an opaque feature blob."
    )

    add_heading(doc, "4.5.4 Q-Learning Based Post Selection", 3)
    add_body_paragraph(
        doc,
        "The QLearningAgent module discretizes each post into a compact state representation using engagement level, recency, text length, hashtag presence, and URL presence. It then evaluates whether a post should be selected or skipped. The agent includes a full temporal-difference update rule and an epsilon-greedy action policy, so the module is structurally consistent with reinforcement-learning literature."
    )
    add_body_paragraph(
        doc,
        "In the current full-pipeline execution, however, the code uses training=False during the selection stage, which means the pipeline behaves as a deterministic selector over the current Q-table state rather than as a continually learning agent. This is a sensible implementation choice for an academic prototype because it keeps the behavior reproducible while still preserving the room for future learning-based upgrades."
    )
    add_body_paragraph(
        doc,
        "The database currently stores the selected posts and their q_value fields, but it does not yet populate q_learning_log during routine execution. That observation is important for the final report because it prevents us from overstating the maturity of the reinforcement-learning component. The better conclusion is that Q-learning has been scaffolded and partially activated, but not yet completed as a full closed-loop training subsystem."
    )

    add_heading(doc, "4.5.5 BERT Contextual Embedding Extraction", 3)
    add_body_paragraph(
        doc,
        "The BERTEncoder service is one of the strongest pieces of the implementation. It lazily loads bert-base-uncased, moves the model to GPU when available, and extracts the CLS token as a 768-dimensional contextual embedding. The embeddings are stored in JSON form so that the values can be persisted, audited, and reused by later training stages."
    )
    add_body_paragraph(
        doc,
        "Unlike older bag-of-words or TF-IDF methods, this representation captures the local and global context of each post. A short sentence about work, study, family, or stress can therefore be encoded in a way that preserves semantic meaning. In the context of personality prediction, that is useful because many signals are subtle and depend on style, not only on keywords."
    )
    add_body_paragraph(
        doc,
        "The orchestrator verifies that each embedding is actually persisted in the database before continuing to later phases. This is a good engineering practice because it prevents a mismatch between in-memory processing and stored artifacts. It also means that the chapter can legitimately discuss the system as a durable pipeline rather than a transient notebook experiment."
    )

    add_heading(doc, "4.5.6 Simplified GAN-Style Augmentation", 3)
    add_body_paragraph(
        doc,
        "The report now describes the augmentation module precisely: it is a simplified GAN-style augmenter rather than a full generative adversarial training loop. The GANAugmenter perturbs embedding vectors with Gaussian noise and creates synthetic text using template-based generation. This approach still serves the research purpose of expanding the training set, but it should not be confused with a fully trained generator-discriminator pair."
    )
    add_body_paragraph(
        doc,
        "Within the orchestrator, the synthetic records are written to the SYNTHETIC_DATA model with a generator version label and a generation confidence score. The stored dataset currently contains 133 synthetic rows. That number is meaningful because it shows that the system is not only predicting but also generating extra training material that can be reused across later runs."
    )
    add_body_paragraph(
        doc,
        "From a methodological standpoint, the simplified augmenter is acceptable for a final-year project because it fulfills the role of data expansion without making claims that the code does not support. The correction is valuable because it protects the credibility of the project during viva questions where a panel member may ask exactly how the synthetic samples were produced."
    )

    add_heading(doc, "4.5.7 Lasso and ElasticNet Regression", 3)
    add_body_paragraph(
        doc,
        "The final prediction stage uses the LassoTrainer class. The trainer standardizes the features, normalizes the labels to the 0-to-1 interval, and trains one model per Big Five trait. Depending on the dataset size, the code uses either LassoCV or ElasticNetCV for a more data-driven regularization choice, and then stores the resulting coefficients, intercepts, and evaluation metrics in the LASSO_MODEL table."
    )
    add_body_paragraph(
        doc,
        "This implementation is more defensible than the phrase 'manual Lasso from scratch' that appeared in the original draft. It is still academically solid because it shows proper understanding of the regression concept, but it is honest about the use of a mature machine-learning library for optimization and validation. The code also exposes feature importance through the non-zero coefficients, which is useful for interpretability discussions."
    )
    add_body_paragraph(
        doc,
        "The stored model statistics show that the trait-specific training behavior is not identical across all OCEAN dimensions. That difference is expected because some traits are easier to learn from language than others. In the current database snapshot, extraversion has the strongest average training R-squared among the trait models, while neuroticism and agreeableness are more difficult. This pattern is consistent with the idea that some personality dimensions are more directly signaled by social media text than others."
    )

    add_heading(doc, "4.5.8 Psychometric Profiles, Insights and Public Demo", 3)
    add_body_paragraph(
        doc,
        "After prediction, the system stores a final PSYCHOMETRIC_PROFILE record for each volunteer. That record holds the predicted OCEAN scores, the error metrics against ground truth, the total posts analyzed, the number of embeddings used, the synthetic data count, the confidence score, and a JSON pipeline summary. This makes the system auditable because every run leaves a compact and structured trail in the database."
    )
    add_body_paragraph(
        doc,
        "The insight engine then converts those predictions into domain-oriented recommendations for education, health and wellbeing, employment, and responsible AI. The public demo uses a lightweight heuristic prediction endpoint, not the full training pipeline, so the report should treat it as a demonstrator rather than as the main research model. This distinction again matters because it keeps the implementation description honest."
    )
    add_body_paragraph(
        doc,
        "On the frontend, the radar chart visualization gives a quick visual summary of the predicted Big Five scores. This is useful for researchers because it turns numerical output into a shape that can be grasped at a glance. The dashboard and reports views then help the researcher move from individual profiles to aggregate study supervision."
    )

    add_heading(doc, "4.6 Database Design and Observed Runtime Statistics", 2)
    add_body_paragraph(
        doc,
        "The database schema is intentionally structured around the life cycle of a volunteer. The VOLUNTEER table anchors the study, BFI_SURVEY stores the ground truth, POST captures the social media content, BERT_EMBEDDING stores contextual features, SYNTHETIC_DATA stores the augmentation outputs, LASSO_MODEL stores the trained regression parameters, and PSYCHOMETRIC_PROFILE stores the final predictions. The Q_LEARNING_LOG table exists as a design artifact even though the current routine pipeline does not yet write episode-by-episode logs into it."
    )
    add_body_paragraph(
        doc,
        "This relational structure is a good fit for the project because it preserves traceability from raw text to final output. It also makes it easier to defend the work in an academic setting because each stage can be inspected independently. The stored values show that the implementation has moved beyond a prototype diagram and into a real database-backed system with measurable execution history."
    )
    runtime_rows = [
        ["Item", "Observed Count"],
        ["Volunteers", str(stats["volunteer"])],
        ["BFI surveys", str(stats["bfi_survey"])],
        ["Posts", str(stats["post"])],
        ["BERT embeddings", str(stats["bert_embedding"])],
        ["Q-learning logs", str(stats["q_learning_log"])],
        ["Synthetic data rows", str(stats["synthetic_data"])],
        ["Lasso models", str(stats["lasso_model"])],
        ["Psychometric profiles", str(stats["psychometric_profile"])],
    ]
    add_table(doc, runtime_rows, col_widths=[2.4, 2.2])
    add_body_paragraph(
        doc,
        f"The current database snapshot shows 13 volunteers, 13 BFI surveys, 254 posts, 100 persisted BERT embeddings, 133 synthetic augmentation records, 40 stored Lasso models, and 8 psychometric profiles. The same snapshot also shows 100 posts with q_value values and 100 posts flagged as embedded, which indicates that the main processing path has executed successfully for a sizeable subset of the study data."
    )

    add_heading(doc, "4.7 Testing and Validation", 2)
    add_body_paragraph(
        doc,
        "Validation of the implementation is not limited to one-off manual inspection. The repository includes a validation script that checks the presence of the eight core database models, the six main machine-learning services, the BFI scorer, the pipeline order, the core views, and the Celery setup. This gives the project a more professional engineering posture because it can be checked systematically rather than only narrated in prose."
    )
    add_body_paragraph(
        doc,
        "At runtime, the pipeline also performs its own safeguards. The BERT stage verifies that embeddings were persisted correctly before continuing. The orchestrator exports cleaned timelines so that the preprocessing outcome can be reviewed. The Lasso stage saves both metrics and model coefficients. These checks reduce the chance that a hidden failure will silently propagate through the rest of the system."
    )
    validation_rows = [
        ["Validation Aspect", "Observed Status"],
        ["Model layer", "All 8 expected models are present in the schema"],
        ["Service layer", "All core services are importable and active"],
        ["BFI scorer", "Reverse scoring and OCEAN calculations operate correctly"],
        ["Pipeline order", "Input -> Q-learning -> BERT -> augmentation -> regression"],
        ["Embedding persistence", "Verified before downstream augmentation"],
        ["Prediction storage", "Profiles are written to the database for later review"],
    ]
    add_table(doc, validation_rows, col_widths=[2.6, 4.1])
    add_body_paragraph(
        doc,
        "One important validation finding is that the q_learning_log table is still empty in the current runtime snapshot. This does not mean the project is broken; it means the learning log has not yet been wired into the main execution flow. That is a limitation worth stating clearly because it is precisely the sort of detail a panel may ask about during defense."
    )

    add_heading(doc, "4.8 Results and Discussion", 2)
    add_body_paragraph(
        doc,
        "The observed results should be discussed carefully and honestly. The system is clearly functional, but its predictive quality varies across volunteers. Across the eight stored psychometric profiles, the mean overall MAE is approximately 0.474, the mean correlation is approximately 0.504, and the mean prediction confidence is approximately 0.772. These values show that the model is extracting signal, but also that the current dataset is still too small for stable large-scale generalization."
    )
    add_body_paragraph(
        doc,
        "Some runs perform much better than others. The best stored profile has a positive correlation above 0.9 and a relatively modest MAE, while another run has a negative correlation and a much weaker R-squared value. This spread is not surprising in a small research dataset where the number of posts differs by volunteer and where the text available for some participants is sparse. It does, however, show that the pipeline is sensitive to data availability, which is precisely why the report should avoid overclaiming robustness."
    )
    add_body_paragraph(
        doc,
        "At the regression-model level, the average training metrics across stored Lasso models are more encouraging. The average training R-squared is highest for extraversion and openness, while agreeableness and neuroticism remain more difficult. That pattern suggests that some personality traits are better represented in linguistic style than others. It also indicates that model interpretability may be more useful for explaining relative trait differences than for claiming near-perfect prediction accuracy."
    )
    trait_rows = [
        ["Trait", "Avg Train MAE", "Avg Train R2", "Avg Samples Used", "Avg Synthetic Used"],
    ]
    for trait, avg_mae, avg_r2, avg_samples, avg_syn in stats["lasso_trait_summary"]:
        trait_rows.append([
            trait.title(),
            f"{avg_mae:.4f}",
            f"{avg_r2:.4f}",
            f"{avg_samples:.2f}",
            f"{avg_syn:.2f}",
        ])
    add_table(doc, trait_rows, col_widths=[1.4, 1.4, 1.2, 1.5, 1.5])
    add_body_paragraph(
        doc,
        "The table above should not be read as proof of production-grade model maturity. Instead, it should be interpreted as evidence that the regression layer is trainable, traceable, and capable of producing a sparse explanation of the embedding space. In a final-year project context, that is valuable because it demonstrates both implementation depth and a realistic understanding of current limitations."
    )
    add_body_paragraph(
        doc,
        "A subtle but important finding is that the live prediction endpoint is intentionally heuristic. It uses text length and word count to produce a quick public demo response, which is useful for showing responsiveness but should not be mistaken for the main research model. The actual research output is the database-backed pipeline that uses BERT, augmentation, and regularized regression on the volunteer records."
    )

    add_heading(doc, "4.9 Chapter Summary", 2)
    add_body_paragraph(
        doc,
        "This chapter has shown the implementation as a practical system with real data, real database records, and real limitations. The code supports BFI-44 import, X timeline acquisition, preprocessing, BERT embeddings, simplified augmentation, Lasso-based prediction, and dashboard interpretation. At the same time, it is important to state that some components are less mature than the original wording suggested. The revised report therefore aligns the documentation with the system that actually exists."
    )


def chapter_five(doc, stats):
    doc.add_page_break()
    add_heading(doc, "CHAPTER FIVE", 1)
    add_heading(doc, "SUMMARY, CONCLUSION AND RECOMMENDATIONS", 1)

    add_heading(doc, "5.1 Introduction", 2)
    add_body_paragraph(
        doc,
        "The preceding chapters described the motivation, background, design, and implementation of the personality prediction system. This final chapter summarizes the work, states the main conclusion drawn from the implementation, and proposes realistic recommendations for improvement. In doing so, the chapter treats the project as an academic software artifact that has been built, tested, corrected, and evaluated rather than as a purely theoretical proposal."
    )
    add_body_paragraph(
        doc,
        "This is important because a final-year project should close the loop between intention and delivery. The system in this study now does that. It accepts volunteers, stores and scores BFI-44 survey data, retrieves social media posts, cleans them, applies a selection mechanism, creates BERT embeddings, augments the data, trains regularized regression models, and presents the results through a dashboard and a live demo interface."
    )

    add_heading(doc, "5.2 Summary of the Study", 2)
    add_body_paragraph(
        doc,
        "The core objective of the study was to build a personality prediction application that could infer the Big Five traits from social media-style text while retaining traceability and interpretability. The final implementation meets that objective at a prototype level. It uses a clear pipeline and a relational database to hold every major artifact, which means that the system does not just compute a prediction; it also records how that prediction was produced."
    )
    add_body_paragraph(
        doc,
        "The system design places special emphasis on four operational domains. In education, the predictions can be turned into learning recommendations. In health and wellbeing, the personality profile can help identify a need for more structure or stress management. In employment, the outputs can support role fit and team composition discussions. In responsible AI, the project demonstrates how a transparent and auditable model can be used in a sensitive decision-support setting without claiming more certainty than the data allows."
    )
    add_body_paragraph(
        doc,
        "The implementation also makes a stronger point that is easy to overlook in a report: a useful project is not only about model accuracy. It is also about data flow, persistence, validation, and reproducibility. By storing the BFI surveys, the cleaned posts, the embeddings, the synthetic samples, the model parameters, and the final profiles, the application provides a complete evidence chain for each prediction."
    )

    add_heading(doc, "5.3 Major Contributions of the Project", 2)
    add_body_paragraph(
        doc,
        "The first contribution is architectural. The project combines psychometric survey data and social media text inside one Django application. That integration matters because it shows how a small research tool can be turned into a usable web system with researchers, volunteers, dashboards, and operational workflows. The modular structure also means that each layer can be improved independently in future work."
    )
    add_body_paragraph(
        doc,
        "The second contribution is methodological. The report now presents a corrected account of the machine-learning stack: a practical post selection mechanism, a BERT embedding service, a simplified augmentation module, and a sparse regression layer. The value here is not in claiming an impossible level of sophistication. It is in showing that the project has been implemented honestly and that each component has a clear role in the pipeline."
    )
    add_body_paragraph(
        doc,
        "The third contribution is traceability. The use of dedicated database models for volunteers, surveys, posts, embeddings, synthetic data, Lasso models, and psychometric profiles gives the system a structure that can be defended in a viva. If a panel member asks where a prediction came from, the answer is not a vague statement. The answer is a path through stored records and reproducible services."
    )
    add_body_paragraph(
        doc,
        "The fourth contribution is usability. The inclusion of a public-facing landing page, a live prediction demo, and a researcher dashboard turns the project from a script into an application. Even where the demo endpoint uses a lightweight heuristic, it still provides a responsive interface that can be shown during demonstration while the deeper research pipeline remains available in the backend."
    )

    add_heading(doc, "5.4 Conclusion", 2)
    add_body_paragraph(
        doc,
        "The study demonstrates that Big Five personality prediction can be implemented as a practical web application using a combination of psychometric scoring, natural language processing, augmentation, and sparse regression. The final system is not perfect, and the report should not present it as if it were. However, it is a real and functioning academic prototype with enough depth to show serious understanding of both the software engineering and the machine-learning sides of the problem."
    )
    add_body_paragraph(
        doc,
        "The most important conclusion from the implementation is that the data matters as much as the algorithm. The database snapshot shows that the system can process a meaningful number of volunteers and posts, but the predictive results also show substantial variance. That variance reminds us that personality prediction from social media text is a difficult problem and that model quality depends heavily on sample size, data diversity, and the quality of the ground truth."
    )
    add_body_paragraph(
        doc,
        "A second conclusion is that transparency should be preferred over overstatement. The corrected report now states that the augmentation module is simplified, that the reinforcement-learning component is scaffolded rather than fully self-learning in routine execution, and that the public demo is heuristic. These clarifications improve the credibility of the work because they allow the document to be defended using the actual codebase rather than an embellished description."
    )
    add_body_paragraph(
        doc,
        "Finally, the project shows that an academic system can still be useful even when it is not fully industrialized. The important thing is that the architecture is coherent, the data path is traceable, the results are measurable, and the limitations are understood. On those grounds, the project is a valid and respectable final-year software engineering and machine-learning study."
    )

    add_heading(doc, "5.5 Limitations of the Study", 2)
    add_body_paragraph(
        doc,
        "The first limitation is dataset size. The current database contains 13 volunteers, 254 posts, and 8 completed psychometric profiles. That is enough for a demonstration, but it is not enough for a fully stable personality model. Small sample sizes make the performance metrics unstable, which is reflected in the spread of correlations and R-squared values across stored profiles."
    )
    add_body_paragraph(
        doc,
        "The second limitation is data heterogeneity. Some volunteers have many more usable posts than others, and this difference affects how much linguistic evidence the model can extract. In practice, personality inference is better when the text is varied and expressive. Sparse or repetitive timelines reduce the signal available to the BERT encoder and the regression layer."
    )
    add_body_paragraph(
        doc,
        "The third limitation is augmentation fidelity. The current augmenter is helpful, but it is not a full adversarial network trained through generator-discriminator competition. It creates useful synthetic samples, yet it does not fully model the statistical complexity of a genuine GAN. The report should therefore treat the augmentation stage as a pragmatic extension rather than as a final solution."
    )
    add_body_paragraph(
        doc,
        "The fourth limitation is reinforcement learning maturity. The Q-learning module is structurally implemented, but the runtime snapshot shows that the logging table is not yet being populated during routine pipeline runs. This means the model selection logic still behaves more like a controlled selector than a continuously learning agent. That is acceptable for an academic prototype, but it should be improved before any stronger claims are made."
    )
    add_body_paragraph(
        doc,
        "The fifth limitation is external validity. The system was tested on the local dataset and the current runtime environment, so its generalization to a much larger, more diverse, or multilingual population is unknown. The model is therefore best treated as a proof of concept that needs more data and more validation before it can be relied on outside the study context."
    )

    add_heading(doc, "5.6 Recommendations", 2)
    recommendation_rows = [
        ["Recommendation", "Reason", "Priority"],
        ["Expand the volunteer sample", "A larger and more diverse dataset will stabilize the regression models and reduce variance.", "High"],
        ["Log Q-learning episodes", "Persisting reward updates will make the selection module truly auditable and trainable.", "High"],
        ["Replace simplified augmentation with a true GAN", "A genuine adversarial model would better match the original research intent.", "Medium"],
        ["Add calibration and external validation", "Better validation will make the reported metrics more defensible.", "High"],
        ["Improve demo prediction path", "The public demo should eventually call the trained pipeline rather than a heuristic fallback.", "Medium"],
        ["Document ethics and consent more clearly", "Personality prediction is sensitive and requires visible governance controls.", "High"],
    ]
    add_table(doc, recommendation_rows, col_widths=[2.5, 3.2, 0.8])
    add_body_paragraph(
        doc,
        "These recommendations are intentionally practical. They do not assume unlimited resources or a completely new research direction. Instead, they focus on the next sensible steps that would make the current system stronger, more transparent, and more defensible in an academic review."
    )

    add_heading(doc, "5.7 Suggestions for Future Work", 2)
    add_body_paragraph(
        doc,
        "Future work should focus on improving the quality and richness of the signal available to the system. One direction is to include profile biography text, pinned posts, and interaction metadata so that the model can learn from more than just timeline content. Another direction is to introduce multimodal signals where ethically appropriate, although that would require a more careful discussion of privacy and consent."
    )
    add_body_paragraph(
        doc,
        "A second direction is model refinement. The current architecture can be upgraded with a genuine reinforcement-learning loop, a true conditional GAN or another modern generative strategy, and more advanced calibration methods. However, such upgrades should only be adopted if they improve measurable performance and do not reduce interpretability in the process."
    )
    add_body_paragraph(
        doc,
        "A third direction is deployment readiness. If the project is to be used by other researchers, then the application should support stronger authentication, role-based permissions, richer logging, backup and recovery, and a clearer API layer for model access. That would turn the current academic prototype into a more durable platform for future research."
    )
    add_body_paragraph(
        doc,
        "A final direction is ethical governance. Personality inference systems can easily be misused if they are presented as infallible. Future versions should therefore include clear consent screens, explanation tools, explicit uncertainty displays, and documented use limitations. Those features will help ensure that the system remains a research tool rather than becoming an opaque screening mechanism."
    )

    add_heading(doc, "5.8 Final Closing Statement", 2)
    add_body_paragraph(
        doc,
        "In summary, the project achieved its central academic goal: it produced a working personality prediction application and a report that now matches the implementation more closely. The system is useful, the write-up is more accurate, and the remaining weaknesses are now clearly identified. That combination makes the project far stronger than a polished but inaccurate report would have been."
    )


def main():
    doc = Document(str(SOURCE_DOCX))
    patch_existing_claims(doc)
    replace_abstract(doc)
    stats = query_db_stats()
    chapter_four(doc, stats)
    chapter_five(doc, stats)
    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
